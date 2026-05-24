from __future__ import annotations

import logging
import re
import shutil
import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from formatter import (
    add_text_with_reference_run,
    clear_paragraph_content,
    clone_paragraph_properties_without_numbering,
    copy_paragraph_format,
    first_text_run,
    split_generated_paragraphs,
)
from heading_detector import HeadingNode


PLACEHOLDER_PATTERN = re.compile(
    r"^\s*(?:"
    r"【?请填写.{0,30}】?|"
    r"【?请.{0,8}填写】?|"
    r"此处填写.*|"
    r"待补充.*|"
    r"XXX+|"
    r"xxx+|"
    r"\.{3,}|"
    r"…+|"
    r"_{3,}|"
    r"-{3,}"
    r")\s*$"
)


@dataclass
class SectionWriteRequest:
    heading: HeadingNode
    content: str
    status: str = "generated"


@dataclass
class SectionWriteResult:
    title: str
    paragraph_index: int
    status: str
    inserted_paragraphs: int
    insertion_mode: str
    reason: str = ""


def write_generated_sections(
    *,
    template_path: Path,
    output_path: Path,
    headings: list[HeadingNode],
    sections: list[SectionWriteRequest],
    logger: Optional[logging.Logger] = None,
) -> list[SectionWriteResult]:
    """Copy template to a temp output and insert generated content in-place."""
    logger = logger or logging.getLogger(__name__)
    temp_path = output_path.with_name(f".{output_path.stem}.{uuid.uuid4().hex}.tmp{output_path.suffix}")

    shutil.copy2(template_path, temp_path)
    document = Document(str(temp_path))
    all_heading_indexes = {heading.paragraph_index for heading in headings}

    results: list[SectionWriteResult] = []
    by_index = {request.heading.paragraph_index: request for request in sections}

    try:
        for paragraph_index in sorted(by_index.keys(), reverse=True):
            request = by_index[paragraph_index]
            result = _insert_section_content(
                document=document,
                request=request,
                heading_indexes=all_heading_indexes,
            )
            logger.info(
                "Inserted section %s at paragraph %s with %s paragraph(s)",
                request.heading.title,
                paragraph_index,
                result.inserted_paragraphs,
            )
            results.append(result)

        document.save(str(temp_path))
        if output_path.exists():
            output_path.unlink()
        shutil.move(str(temp_path), str(output_path))
    except Exception:
        logger.exception("Failed to write generated DOCX")
        if temp_path.exists():
            temp_path.unlink()
        raise

    return sorted(results, key=lambda item: item.paragraph_index)


def _insert_section_content(
    *,
    document: DocumentObject,
    request: SectionWriteRequest,
    heading_indexes: set[int],
) -> SectionWriteResult:
    heading = request.heading
    paragraphs = document.paragraphs
    if heading.paragraph_index >= len(paragraphs):
        return SectionWriteResult(
            title=heading.title,
            paragraph_index=heading.paragraph_index,
            status="failed",
            inserted_paragraphs=0,
            insertion_mode="none",
            reason="Heading paragraph index no longer exists in output document.",
        )

    body_paragraphs = split_generated_paragraphs(request.content)
    if not body_paragraphs:
        return SectionWriteResult(
            title=heading.title,
            paragraph_index=heading.paragraph_index,
            status="skipped",
            inserted_paragraphs=0,
            insertion_mode="none",
            reason="Generated content is empty after cleanup.",
        )

    anchor = paragraphs[heading.paragraph_index]
    placeholder = _find_placeholder_after_heading(
        paragraphs=paragraphs,
        heading=heading,
        heading_indexes=heading_indexes,
    )
    reference = _find_reference_body_paragraph(
        document=document,
        heading=heading,
        heading_indexes=heading_indexes,
        placeholder=placeholder,
    )
    reference_run = first_text_run(reference)

    if placeholder is not None:
        _rewrite_paragraph(placeholder, body_paragraphs[0], reference, reference_run)
        current = placeholder
        inserted = 1
        for text in body_paragraphs[1:]:
            current = _insert_paragraph_after(current, text, reference, reference_run)
            inserted += 1
        return SectionWriteResult(
            title=heading.title,
            paragraph_index=heading.paragraph_index,
            status="inserted",
            inserted_paragraphs=inserted,
            insertion_mode="replace_placeholder",
        )

    current = anchor
    inserted = 0
    for text in body_paragraphs:
        current = _insert_paragraph_after(current, text, reference, reference_run)
        inserted += 1

    return SectionWriteResult(
        title=heading.title,
        paragraph_index=heading.paragraph_index,
        status="inserted",
        inserted_paragraphs=inserted,
        insertion_mode="after_heading",
    )


def _insert_paragraph_after(
    anchor: Paragraph,
    text: str,
    reference: Optional[Paragraph],
    reference_run,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    _apply_reference_format(paragraph, reference)
    add_text_with_reference_run(paragraph, text, reference_run)
    return paragraph


def _rewrite_paragraph(
    paragraph: Paragraph,
    text: str,
    reference: Optional[Paragraph],
    reference_run,
) -> None:
    _apply_reference_format(paragraph, reference)
    clear_paragraph_content(paragraph)
    add_text_with_reference_run(paragraph, text, reference_run)


def _apply_reference_format(paragraph: Paragraph, reference: Optional[Paragraph]) -> None:
    copy_paragraph_format(paragraph, reference)
    clone_paragraph_properties_without_numbering(paragraph, reference)


def _find_placeholder_after_heading(
    *,
    paragraphs: list[Paragraph],
    heading: HeadingNode,
    heading_indexes: set[int],
) -> Optional[Paragraph]:
    upper = min(heading.end_index, len(paragraphs) - 1)
    for index in range(heading.paragraph_index + 1, upper + 1):
        if index in heading_indexes:
            return None
        text = paragraphs[index].text.strip()
        if not text:
            continue
        return paragraphs[index] if PLACEHOLDER_PATTERN.match(text) else None
    return None


def _find_reference_body_paragraph(
    *,
    document: DocumentObject,
    heading: HeadingNode,
    heading_indexes: set[int],
    placeholder: Optional[Paragraph],
) -> Optional[Paragraph]:
    paragraphs = document.paragraphs
    upper = min(heading.end_index, len(paragraphs) - 1)
    if placeholder is not None:
        return placeholder

    for index in range(heading.paragraph_index + 1, upper + 1):
        if index in heading_indexes:
            continue
        paragraph = paragraphs[index]
        if paragraph.text.strip():
            return paragraph

    for index, paragraph in enumerate(paragraphs):
        if index in heading_indexes:
            continue
        if paragraph.text.strip():
            return paragraph

    try:
        style = document.styles["Normal"]
    except Exception:
        return None
    synthetic = document.add_paragraph()
    synthetic.style = style
    synthetic._element.getparent().remove(synthetic._element)
    return synthetic
