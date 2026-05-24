from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

try:
    from docx import Document
    from docx.document import Document as DocumentObject
except ImportError as exc:  # pragma: no cover
    raise RuntimeError("python-docx is required. Install dependencies with pip install -r requirements.txt") from exc


@dataclass(frozen=True)
class DocumentFacts:
    paragraph_count: int
    table_count: int
    inline_shape_count: int
    section_count: int
    has_toc_like_text: bool

    def to_prompt_text(self) -> str:
        return (
            f"Paragraphs: {self.paragraph_count}; tables: {self.table_count}; "
            f"inline images/shapes: {self.inline_shape_count}; sections: {self.section_count}; "
            f"TOC-like text detected: {'yes' if self.has_toc_like_text else 'no'}."
        )


def load_docx(path: Path) -> DocumentObject:
    if not path.exists():
        raise FileNotFoundError(f"Template not found: {path}")
    if path.suffix.lower() != ".docx":
        raise ValueError(f"Template must be a .docx file: {path}")
    return Document(str(path))


def collect_document_facts(document: DocumentObject) -> DocumentFacts:
    texts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    has_toc_like_text = any(text in {"目录", "Contents", "CONTENTS"} for text in texts[:20])
    return DocumentFacts(
        paragraph_count=len(document.paragraphs),
        table_count=len(document.tables),
        inline_shape_count=len(document.inline_shapes),
        section_count=len(document.sections),
        has_toc_like_text=has_toc_like_text,
    )
